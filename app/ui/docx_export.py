"""Build .docx downloads for briefings and API setup guide."""

from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import RGBColor


def _save_doc(doc: Document) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_meta_line(doc: Document, label: str, value: str) -> None:
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    para.add_run(value)


def _add_markdownish_body(doc: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            plain = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            doc.add_paragraph(plain)


def build_briefing_docx(result: dict) -> bytes:
    """Export a briefing result as a formatted Word document."""
    doc = Document()
    title = doc.add_heading("World Cup Match Briefing", level=0)
    title.runs[0].font.color.rgb = RGBColor(0xBE, 0x12, 0x3C)

    _add_meta_line(doc, "Generated", datetime.now().strftime("%Y-%m-%d %H:%M"))
    _add_meta_line(doc, "Query", result.get("query", ""))
    if result.get("team_name"):
        _add_meta_line(doc, "Focus team", result["team_name"])
    if result.get("team_name_corrected") and result.get("team_name_input"):
        _add_meta_line(doc, "Resolved from", result["team_name_input"])
    if result.get("next_match"):
        _add_meta_line(doc, "Next match", result["next_match"])

    doc.add_paragraph()
    doc.add_heading("Briefing", level=1)
    _add_markdownish_body(doc, result.get("briefing", ""))
    return _save_doc(doc)


def build_setup_guide_docx() -> bytes:
    """Export API key setup instructions as a Word document."""
    doc = Document()
    heading = doc.add_heading("WorldCup Analyst — API Keys Setup Guide", level=0)
    heading.runs[0].font.color.rgb = RGBColor(0xBE, 0x12, 0x3C)

    intro = doc.add_paragraph()
    intro.add_run(
        "Keys are entered in the app sidebar and stay in your browser session only. "
        "They are never stored on our servers."
    )

    sections = [
        (
            "1. Groq API Key (required)",
            "Powers the AI agents and briefing synthesis.",
            [
                "Go to https://console.groq.com",
                "Sign up or log in",
                'Open "API Keys" in the left menu',
                'Click "Create API Key" (e.g. worldcup-analyst)',
                "Copy the key (starts with gsk_) and paste it into the Groq field",
            ],
        ),
        (
            "2. Football Data Token (required)",
            "Provides live World Cup fixtures, standings, form, and scorers.",
            [
                "Go to https://www.football-data.org/client/register",
                "Register for a free account",
                "Log in and open your dashboard",
                "Copy your X-Auth-Token",
                "Paste it into the Football Data Token field",
            ],
        ),
        (
            "3. Tavily API Key (optional)",
            "Enables richer web news search. Without it, RSS feeds are used.",
            [
                "Go to https://app.tavily.com/home",
                "Sign up for a free account",
                "Create an API key in your dashboard",
                "Copy the key (starts with tvly-) and paste it into the Tavily field",
            ],
        ),
    ]

    for title, summary, steps in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(summary)
        for step in steps:
            doc.add_paragraph(step, style="List Number")

    doc.add_heading("Troubleshooting", level=1)
    tips = [
        '"Please add Groq / Football Data" → both required keys must be filled in',
        "Briefing fails with HTTP 403 → check your Football Data token",
        "Briefing fails with auth error → regenerate your Groq key",
        "No news results → add a Tavily key, or rely on RSS-only mode",
    ]
    for tip in tips:
        doc.add_paragraph(tip, style="List Bullet")

    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run("Football Data free tier allows 10 requests/minute. The app caches responses.")
    return _save_doc(doc)